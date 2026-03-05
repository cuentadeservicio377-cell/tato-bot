"""
data/parse_docx.py — Parsea los DOCX de control de Tato a dicts estructurados.

Funciones:
  parse_control_expedientes(path) → list[dict]
  parse_pendientes(path)          → list[dict]
  merge_pendientes_into_expedientes(expedientes, pendientes) → list[dict]
"""

import re
import uuid
from datetime import datetime
from docx import Document

# ─── Lookup: código juzgado → nombre completo ───────────────────────────────

JUZGADO_NOMBRES = {
    "1M": "Juzgado Primero De Lo Mercantil",
    "2M": "Juzgado Segundo De Lo Mercantil",
    "3M": "Juzgado Tercero De Lo Mercantil",
    "4M": "Juzgado Cuarto De Lo Mercantil",
    "5M": "Juzgado Quinto De Lo Mercantil",
    "6M": "Juzgado Sexto De Lo Mercantil",
    "7M": "Juzgado Séptimo De Lo Mercantil",
    "8M": "Juzgado Octavo De Lo Mercantil",
    "9M": "Juzgado Noveno De Lo Mercantil",
    "10M": "Juzgado Décimo De Lo Mercantil",
    "11M": "Juzgado Décimo Primero De Lo Mercantil",
    "12M": "Juzgado Décimo Segundo De Lo Mercantil",
    "13M": "Juzgado Décimo Tercero De Lo Mercantil",
    "14M": "Juzgado Décimo Cuarto De Lo Mercantil",
    "17M": "Juzgado Décimo Séptimo De Lo Mercantil",
    "12 MO": "Juzgado Décimo Segundo en Materia Oral Mercantil",
    "13O-M": "Juzgado Décimo Tercero Oral Mercantil",
    "14OM": "Juzgado Décimo Cuarto en Materia Oral Mercantil",
    "6OM": "Juzgado Sexto en Materia Oral Mercantil",
    "18MO": "Juzgado Décimo Octavo en Materia Oral Mercantil",
    "1F": "Juzgado Primero De Lo Familiar",
    "2F": "Juzgado Segundo De Lo Familiar",
    "3F": "Juzgado Tercero De Lo Familiar",
    "4F": "Juzgado Cuarto De Lo Familiar",
    "5F": "Juzgado Quinto De Lo Familiar",
    "6F": "Juzgado Sexto De Lo Familiar",
    "7F": "Juzgado Séptimo De Lo Familiar",
    "8F": "Juzgado Octavo De Lo Familiar",
    "9F": "Juzgado Noveno De Lo Familiar",
    "10F": "Juzgado Décimo De Lo Familiar",
    "11F": "Juzgado Décimo Primero De Lo Familiar",
    "12F": "Juzgado Décimo Segundo De Lo Familiar",
    "13F": "Juzgado Décimo Tercero De Lo Familiar",
    "14F": "Juzgado Décimo Cuarto De Lo Familiar",
    "15F": "Juzgado Décimo Quinto De Lo Familiar",
    "16F": "Juzgado Décimo Sexto De Lo Familiar",
    "17F": "Juzgado Décimo Séptimo De Lo Familiar",
    "21F": "Juzgado Vigésimo Primero De Lo Familiar",
    "1C": "Juzgado Primero De Lo Civil",
    "2C": "Juzgado Segundo De Lo Civil",
    "3C": "Juzgado Tercero De Lo Civil",
    "4C": "Juzgado Cuarto De Lo Civil",
    "5C": "Juzgado Quinto De Lo Civil",
    "6C": "Juzgado Sexto De Lo Civil",
    "7C": "Juzgado Séptimo De Lo Civil",
    "8C": "Juzgado Octavo De Lo Civil",
    "9C": "Juzgado Noveno De Lo Civil",
    "10C": "Juzgado Décimo De Lo Civil",
    "11C": "Juzgado Décimo Primero De Lo Civil",
    "12C": "Juzgado Décimo Segundo De Lo Civil",
    "13C": "Juzgado Décimo Tercero De Lo Civil",
    "2D-Civil": "Juzgado Segundo De Distrito En Materia Civil",
    "2DC": "Juzgado Segundo De Distrito En Materia Civil",
    "7DC": "Juzgado Séptimo De Distrito En Materia Civil",
    "19D": "Juzgado Décimo Noveno De Distrito",
    "1D-P": "Juzgado Primero De Distrito En Materia Penal",
    "1D Trab": "Juzgado Primero De Distrito En Materia Laboral",
    "1S": "Primera Sala",
    "II-S": "Segunda Sala",
    "III-S": "Tercera Sala",
    "3S": "Tercera Sala",
    "IV-S": "Cuarta Sala",
    "IV SALA": "Cuarta Sala",
    "V-S": "Quinta Sala",
    "VI-S": "Sexta Sala",
    "VI sala": "Sexta Sala",
    "I sala": "Primera Sala",
    "9S": "Novena Sala",
    "TCC": "Tribunal Colegiado Civil",
    "4TC C": "Cuarto Tribunal Colegiado En Materia Civil",
    "12JE": "Juzgado Décimo Segundo en Materia de Extinción",
    "IJA": "Instituto de Justicia Alternativa",
}

# Keywords para inferir estado
_CADUCIDAD_KW = {"CADUCIDAD", "CADUCADO"}
_TERMINADO_KW = {"AL KILO", "JUICIO CONCLUIDO", "ARCHIVESE CONCLUIDO",
                  "BAJA SIN CARPETA", "SENTENCIA FIRME", "CUMPLIDA SENTENCIA",
                  "EXPEDIENTE TERMINADO"}


def _inferir_estado(observaciones: str, domicilio: str) -> str:
    text = (observaciones + " " + domicilio).upper()
    for kw in _CADUCIDAD_KW:
        if kw in text:
            return "caducidad"
    for kw in _TERMINADO_KW:
        if kw in text:
            return "terminado"
    return "activo"


def _normalizar_codigo(raw: str) -> str:
    """Extrae el primer código de juzgado de un campo que puede tener varios."""
    if not raw:
        return ""
    # Tomar solo el primer token antes de espacio múltiple, coma, etc.
    first = re.split(r"\s{2,}|,|/|\\n", raw)[0].strip()
    return first


def _extraer_actor(partes: str) -> str:
    """Extrae la parte actora (antes del 'vs' / 'VS' / 'Vs')."""
    if not partes:
        return ""
    sep = re.split(r"\s+[Vv][Ss]\.?\s+", partes, maxsplit=1)
    return sep[0].strip() if sep else partes.strip()


def parse_control_expedientes(docx_path: str) -> list:
    """
    Parsea la tabla del DOCX CONTROL EXPEDIENTES.
    Retorna lista de dicts con schema completo de expediente.
    Incluye activos, terminados y caducidad (estado en campo 'estado').
    """
    doc = Document(docx_path)
    if not doc.tables:
        return []

    table = doc.tables[0]
    rows = table.rows
    result = []

    for row in rows[1:]:  # skip header
        cells = [c.text.strip().replace("\r", " ").replace("\n", " ") for c in row.cells]
        # Ensure 7 cols
        while len(cells) < 7:
            cells.append("")

        numero_interno = cells[0].strip()
        juz_raw        = cells[1].strip()
        exp_raw        = cells[2].strip()
        partes         = cells[3].strip()
        monto          = cells[4].strip()
        domicilio      = cells[5].strip()
        observaciones  = cells[6].strip()

        # Skip completamente vacío o sin expediente y sin partes
        if not exp_raw and not partes:
            continue
        # Skip filas de notas sueltas (sin JUZ ni EXP pero con texto en PARTES)
        if not exp_raw and not juz_raw:
            continue

        juzgado_codigo = _normalizar_codigo(juz_raw)
        juzgado_nombre = JUZGADO_NOMBRES.get(juzgado_codigo, juzgado_codigo)

        # Si el código tiene múltiples juzgados (ej: "4M 8S"), usar el primero
        if not juzgado_nombre or juzgado_nombre == juzgado_codigo:
            # Intenta normalizar a mayúsculas
            juzgado_nombre = JUZGADO_NOMBRES.get(juzgado_codigo.upper(), juzgado_codigo)

        # Expediente: puede ser "1574/19 124/20" → tomar el primero
        numero = re.split(r"\s{2,}|\n", exp_raw)[0].strip()

        cliente = _extraer_actor(partes)
        estado  = _inferir_estado(observaciones, domicilio)

        expediente = {
            "id": str(uuid.uuid4()),
            "numero_interno": numero_interno,
            "numero": numero,
            "juzgado": juzgado_nombre,
            "juzgado_codigo": juzgado_codigo,
            "partes": partes,
            "cliente": cliente,
            "tipo": "",
            "monto": monto,
            "domicilio_demandado": domicilio,
            "estado": estado,
            "etapa": "",
            "ultimo_acuerdo": "",
            "ultimo_acuerdo_texto": "",
            "proximo_paso": "",
            "proximo_termino": "",
            "termino_fatal": False,
            "notas": observaciones,
            "ultima_actualizacion": datetime.now().isoformat(),
            "sheets_row": None,
        }
        result.append(expediente)

    return result


def _separar_acuerdo_nota(texto: str):
    """
    Separa el acuerdo oficial de la nota de Tato.
    Patrón: acuerdo oficial (ALL CAPS/formal) seguido de " . " o ". " + nota informal.
    Retorna (acuerdo, nota).
    """
    if not texto:
        return "", ""

    # Busca separador: ". [minúsculas/nota]"
    # La nota de Tato suele ser texto en minúsculas o mixed case después del acuerdo
    # Patrón: último ". " seguido de texto que empieza en minúscula o con palabras de acción
    match = re.search(r"\.\s+([a-z\-–])", texto)
    if match:
        split_idx = match.start()
        acuerdo = texto[:split_idx].strip().rstrip(".")
        nota    = texto[match.start() + 2:].strip()
        return acuerdo, nota

    # Alternativa: " . " (espacio-punto-espacio seguido de texto informal)
    parts = texto.split(" . ")
    if len(parts) >= 2:
        # La última parte puede ser la nota
        last = parts[-1].strip()
        # Si tiene texto informal (mayúsculas mezcladas o específicas)
        acuerdo = " . ".join(parts[:-1]).strip()
        return acuerdo, last

    return texto.strip(), ""


def _extraer_codigo_pendientes(juzgado_str: str):
    """
    Del campo juzgado en PENDIENTES (ej: '5M Juzgado Quinto De Lo Mercantil'),
    extrae el código y el nombre.
    """
    if not juzgado_str:
        return "", ""

    # Intenta extraer código al inicio (ej: "5M", "3F", "13F", "2D-Civil", "III-S")
    match = re.match(r"^([0-9A-Za-zÑñ\-]+)\s+(.+)", juzgado_str.strip())
    if match:
        codigo = match.group(1).strip()
        nombre = match.group(2).strip()
        return codigo, nombre

    return juzgado_str.strip(), juzgado_str.strip()


def parse_pendientes(docx_path: str) -> list:
    """
    Parsea la tabla del DOCX JUZGADOS PENDIENTES.
    Retorna lista de dicts con:
      juzgado, juzgado_codigo, numero, ultimo_acuerdo, ultimo_acuerdo_texto, proximo_paso
    """
    doc = Document(docx_path)
    if not doc.tables:
        return []

    table = doc.tables[0]
    result = []

    for row in table.rows:
        cells = [c.text.strip().replace("\r", " ").replace("\n", " ") for c in row.cells]
        while len(cells) < 4:
            cells.append("")

        juzgado_raw  = cells[0].strip()
        exp_raw      = cells[1].strip()
        fecha_raw    = cells[2].strip()
        acuerdo_raw  = cells[3].strip()

        if not exp_raw:
            continue

        codigo, nombre = _extraer_codigo_pendientes(juzgado_raw)
        if not nombre:
            nombre = JUZGADO_NOMBRES.get(codigo, codigo)

        # Fecha: "2026-02-13 09:27:55" → solo fecha
        fecha = fecha_raw.split(" ")[0] if fecha_raw else ""

        acuerdo, nota = _separar_acuerdo_nota(acuerdo_raw)

        result.append({
            "numero": exp_raw.strip(),
            "juzgado": nombre,
            "juzgado_codigo": codigo,
            "ultimo_acuerdo": fecha,
            "ultimo_acuerdo_texto": acuerdo,
            "proximo_paso": nota,
        })

    return result


def merge_pendientes_into_expedientes(expedientes: list, pendientes: list) -> list:
    """
    Cross-referencia pendientes con expedientes por numero de expediente.
    Actualiza: ultimo_acuerdo, ultimo_acuerdo_texto, proximo_paso.
    Si la fecha del pendiente es más reciente, reemplaza.
    """
    # Índice de expedientes por número (normalizado)
    idx = {}
    for i, exp in enumerate(expedientes):
        key = exp.get("numero", "").strip()
        if key:
            idx[key] = i

    for p in pendientes:
        key = p.get("numero", "").strip()
        if key not in idx:
            continue
        exp = expedientes[idx[key]]

        # Solo actualizar si el pendiente tiene info más reciente
        fecha_actual    = exp.get("ultimo_acuerdo", "")
        fecha_pendiente = p.get("ultimo_acuerdo", "")

        if not fecha_actual or (fecha_pendiente and fecha_pendiente >= fecha_actual):
            if p.get("ultimo_acuerdo"):
                exp["ultimo_acuerdo"] = p["ultimo_acuerdo"]
            if p.get("ultimo_acuerdo_texto"):
                exp["ultimo_acuerdo_texto"] = p["ultimo_acuerdo_texto"]
            if p.get("proximo_paso"):
                exp["proximo_paso"] = p["proximo_paso"]
            # Si el juzgado del pendiente tiene nombre completo, preferirlo
            if p.get("juzgado") and len(p["juzgado"]) > len(exp.get("juzgado", "")):
                exp["juzgado"] = p["juzgado"]

    return expedientes
